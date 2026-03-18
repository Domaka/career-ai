from io import BytesIO
from types import SimpleNamespace
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from docx import Document
from rest_framework.test import APITestCase

from .models import CvAnalysis, CvExtract

User = get_user_model()


class AuthApiTests(APITestCase):
	def test_register_login_logout(self):
		register = self.client.post(
			reverse("auth-register"),
			{"username": "dom", "password": "pass1234", "email": "dom@example.com"},
			format="json",
		)
		self.assertEqual(register.status_code, 201)
		self.assertIn("token", register.data)

		login = self.client.post(
			reverse("auth-login"),
			{"username": "dom", "password": "pass1234"},
			format="json",
		)
		self.assertEqual(login.status_code, 200)
		token = login.data["token"]

		me = self.client.get(reverse("auth-me"), HTTP_AUTHORIZATION=f"Token {token}")
		self.assertEqual(me.status_code, 200)
		self.assertEqual(me.data["username"], "dom")

		logout = self.client.post(reverse("auth-logout"), HTTP_AUTHORIZATION=f"Token {token}")
		self.assertEqual(logout.status_code, 200)

		me_after_logout = self.client.get(reverse("auth-me"), HTTP_AUTHORIZATION=f"Token {token}")
		self.assertEqual(me_after_logout.status_code, 401)


class CvExtractApiTests(APITestCase):
	def setUp(self):
		self.user = User.objects.create_user(username="tester", password="pass1234")
		self.client.force_authenticate(user=self.user)

	def test_cv_extract_requires_file(self):
		response = self.client.post(reverse("cv-extract"), {"profile_id": 12}, format="multipart")
		self.assertEqual(response.status_code, 400)
		self.assertEqual(response.data["error"], "cv_file is required")

	def test_cv_extract_requires_profile_id(self):
		cv_file = SimpleUploadedFile("sample.docx", b"fake-content")
		response = self.client.post(reverse("cv-extract"), {"cv_file": cv_file}, format="multipart")
		self.assertEqual(response.status_code, 400)
		self.assertEqual(response.data["error"], "profile_id is required")

	@patch("api.views.extract_cv_intelligence")
	def test_cv_extract_success(self, mock_extract):
		mock_extract.return_value = SimpleNamespace(
			raw_text="Raw CV",
			clean_text="Clean CV",
			page_count=1,
			extraction_mode="heuristic_with_gemini_review",
			llm_structured_cv={"summary": "LLM output"},
			comparison={"llm_available": True, "provider": "gemini"},
			derived_metrics={
				"total_experience_months": 24,
				"total_projects": 2,
				"projects_with_metrics": 1,
				"unique_skill_count": 10,
				"leadership_signal_count": 1,
			},
			structured_cv={
				"summary": "Python backend developer",
				"skills": {
					"technical": ["Python"],
					"frameworks": ["Django"],
					"databases": ["Postgres"],
					"devops": ["Docker"],
					"tools": ["Git"],
					"soft_skills": ["Collaboration"],
				},
				"experience": [],
				"projects": [],
				"education": [],
				"certifications": [],
				"career_level_estimate": "mid",
				"core_strengths": [],
				"potential_weaknesses": [],
				"personality_indicators": [],
				"career_trajectory_analysis": "Mid-level progression",
				"cv_weaknesses": [],
				"missing_signals": [],
				"possible_missed_information": [],
				"extractor_learning_insights": [],
				"confidence_score": 86,
			},
		)

		cv_file = SimpleUploadedFile(
			"candidate.docx",
			b"fake-content",
			content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		)
		response = self.client.post(
			reverse("cv-extract"),
			{"profile_id": 99, "target_role": "Backend Django Developer", "cv_file": cv_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, 201)
		self.assertEqual(response.data["profile_id"], 99)
		self.assertEqual(response.data["career_level_estimate"], "mid")
		self.assertEqual(response.data["extraction_mode"], "heuristic_with_gemini_review")
		self.assertEqual(response.data["comparison"]["provider"], "gemini")
		self.assertIn("analysis", response.data)
		self.assertIn("summary", response.data["analysis"])
		self.assertIn("structured_cv", response.data)
		self.assertEqual(CvExtract.objects.count(), 1)
		self.assertEqual(CvAnalysis.objects.count(), 1)

	def test_cv_extract_end_to_end_docx(self):
		document = Document()
		document.add_paragraph("Professional Summary")
		document.add_paragraph("Backend developer focused on APIs and performance.")
		document.add_paragraph("Skills")
		document.add_paragraph("Python, Django, DRF, Postgres, Docker, Git, Communication")
		document.add_paragraph("Work Experience")
		document.add_paragraph("Backend Engineer at Acme | Jan 2022 - Present")
		document.add_paragraph("Led a team of 3 engineers to modernize backend services")
		document.add_paragraph("Improved response times by 35%")
		document.add_paragraph("Projects")
		document.add_paragraph("Career AI")
		document.add_paragraph("Built a CV intelligence API. https://github.com/example/career-ai")

		buffer = BytesIO()
		document.save(buffer)
		temp_file = SimpleUploadedFile(
			"sample.docx",
			buffer.getvalue(),
			content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		)

		response = self.client.post(
			reverse("cv-extract"),
			{"profile_id": 42, "target_role": "Backend Django Developer", "cv_file": temp_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, 201)
		self.assertEqual(response.data["profile_id"], 42)
		self.assertIn("structured_cv", response.data)
		self.assertIn("skills", response.data["structured_cv"])
		self.assertEqual(response.data["extraction_mode"], "heuristic")
		self.assertFalse(response.data["comparison"]["llm_available"])
		self.assertIn("analysis", response.data)
		self.assertIn("strengths", response.data["analysis"])
		self.assertGreaterEqual(response.data["confidence_score"], 0)
		self.assertEqual(CvAnalysis.objects.count(), 1)


class CvExtractAuthRequiredTests(APITestCase):
	def test_cv_extract_requires_authentication(self):
		response = self.client.post(reverse("cv-extract"), {"profile_id": 12}, format="multipart")
		self.assertEqual(response.status_code, 401)
